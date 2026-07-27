# src/extraction/docx_extractor.py

import docx
from src.utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_docx(file_path: str) -> dict:
    logger.info(f"Attempting to extract text from DOCX: {file_path}")
    try:
        document = docx.Document(file_path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        table_text = []
        for table in document.tables:
            for row in table.rows:
                table_text.extend(cell.text for cell in row.cells)
        full_text = "\n".join(paragraphs + table_text)
        metadata = {"filename": file_path, "char_count": len(full_text)}
        logger.debug(
            f"Successfully extracted {len(full_text)} characters from {file_path}"
        )
        return {"text": full_text, "metadata": metadata}
    except Exception as e:
        logger.error(f"Failed to read DOCX {file_path}: {e}", exc_info=True)
        return {
            "text": "",
            "metadata": {"filename": file_path, "char_count": 0, "error": str(e)},
        }
